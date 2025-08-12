import cv2
import numpy as np
from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import io
from datetime import datetime

from analysis.models import ForgeryType



def compare_signatures(path_img1: str, path_img2: str) -> str:
    """
    Compara duas imagens (assinaturas) e retorna um resultado textual básico de similaridade.
    """
    img1 = cv2.imread(path_img1, cv2.IMREAD_GRAYSCALE)
    img2 = cv2.imread(path_img2, cv2.IMREAD_GRAYSCALE)

    if img1 is None or img2 is None:
        return "Uma ou ambas as imagens não puderam ser carregadas."

    img1 = cv2.resize(img1, (300, 150))
    img2 = cv2.resize(img2, (300, 150))

    diff = cv2.absdiff(img1, img2)
    non_zero_count = np.count_nonzero(diff)

    similarity = 1 - non_zero_count / (300 * 150)
    similarity_pct = similarity * 100

    if similarity_pct > 80:
        return f"Alto nível de similaridade detectado ({similarity_pct:.2f}%). Assinaturas muito semelhantes."
    elif similarity_pct > 50:
        return f"Similaridade moderada ({similarity_pct:.2f}%). Recomendado revisão manual."
    else:
        return f"Baixa similaridade ({similarity_pct:.2f}%). Possível divergência ou falsificação."


def generate_case_report_docx(case, analyses, user, output_path):
    """
    Gera um relatório DOCX de um caso com base nas análises.
    """
    doc = DocxDocument()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    doc.add_heading(f'Laudo Técnico – Caso #{case.id}', 0)
    doc.add_paragraph(f'Descrição do Caso: {case.description}')
    doc.add_paragraph(f'Perito: {user.get_full_name()} ({user.email})')
    doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')

    doc.add_heading('Análises Realizadas:', level=1)
    for idx, analysis in enumerate(analyses, 1):
        doc.add_paragraph(f"{idx}. Observação: {analysis.observation}")
        if analysis.automatic_result:
            doc.add_paragraph(f"Resultado Automático: {analysis.automatic_result}")
        if analysis.forgery_type:
            doc.add_paragraph(f"Tipo de Falsificação: {analysis.forgery_type.name}")

    doc.save(output_path)


def generate_case_report_pdf(case, analyses, user):
    """
    Gera um relatório PDF de um caso com base nas análises.
    Retorna o conteúdo PDF em bytes.
    """
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    y = height - 50
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, y, f"Laudo Técnico – Caso #{case.id}")

    p.setFont("Helvetica", 12)
    y -= 30
    p.drawString(50, y, f"Descrição do Caso: {case.description}")
    y -= 20
    p.drawString(50, y, f"Perito: {user.get_full_name()} ({user.email})")
    y -= 20
    p.drawString(50, y, f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    y -= 40
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, y, "Análises Realizadas:")

    p.setFont("Helvetica", 12)
    for idx, analysis in enumerate(analyses, 1):
        y -= 20
        p.drawString(60, y, f"{idx}. Observação: {analysis.observation}")
        if analysis.automatic_result:
            y -= 15
            p.drawString(70, y, f"Resultado Automático: {analysis.automatic_result}")
        if analysis.forgery_type:
            y -= 15
            p.drawString(70, y, f"Tipo de Falsificação: {analysis.forgery_type.name}")

        if y < 100:  # Nova página
            p.showPage()
            y = height - 50
            p.setFont("Helvetica", 12)

    p.showPage()
    p.save()
    pdf = buffer.getvalue()
    buffer.close()
    return pdf


def simulate_ai_analysis(image_path: str) -> str:
    """
    Simula análise automática de uma imagem de assinatura com base em regras fictícias.
    """
    # Aqui podemos no futuro integrar com IA real
    # Por enquanto, apenas retorna texto fictício baseado no nome do arquivo
    if "falso" in image_path.lower():
        return "Alta probabilidade de falsificação detectada pela análise automática."
    elif "suspeito" in image_path.lower():
        return "Indícios de falsificação. Recomendado análise manual detalhada."
    else:
        return "Assinatura consistente com padrões conhecidos. Nenhum indício forte de falsificação."


def get_forgery_patterns():
    """
    Retorna uma lista fictícia de padrões de falsificação conhecidos.
    """
    return [
        "Assinatura tremida",
        "Traços interrompidos",
        "Pressão irregular na escrita",
        "Tamanho inconsistente das letras",
        "Alterações visíveis em partes específicas"
    ]

def get_forgery_types():
    """
    Retorna todos os tipos de falsificação cadastrados no banco.
    """
    return ForgeryType.objects.all()