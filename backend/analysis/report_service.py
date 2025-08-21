import os
import io
import logging
from django.conf import settings
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

logger = logging.getLogger(__name__)


def generate_pdf(case, include_signature=False):
    """
    Gera um PDF do relatório do caso, com ou sem assinatura do perito.
    Retorna o caminho do arquivo.
    """
    reports_dir = os.path.join(settings.MEDIA_ROOT, "reports")
    os.makedirs(reports_dir, exist_ok=True)

    pdf_path = os.path.join(reports_dir, f"case_{case.id}_report.pdf")

    try:
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.drawString(100, 750, f"Relatório do Caso {case.id}")
        c.drawString(100, 730, f"Cliente: {case.cliente_nome if hasattr(case, 'cliente_nome') else 'N/A'}")
        c.drawString(100, 710, f"Descrição: {case.description}")

        if include_signature:
            c.drawString(100, 690, "Assinado digitalmente pelo perito responsável")

        c.save()
        return pdf_path
    except Exception as e:
        logger.error(f"Erro ao gerar PDF: {e}")
        return None


def generate_docx(case, include_signature=False):
    """
    Gera um DOCX do relatório do caso.
    Retorna o caminho do arquivo.
    """
    reports_dir = os.path.join(settings.MEDIA_ROOT, "reports")
    os.makedirs(reports_dir, exist_ok=True)

    docx_path = os.path.join(reports_dir, f"case_{case.id}_report.docx")

    try:
        doc = Document()
        doc.add_heading(f"Relatório do Caso {case.id}", 0)
        doc.add_paragraph(f"Cliente: {case.cliente_nome if hasattr(case, 'cliente_nome') else 'N/A'}")
        doc.add_paragraph(f"Descrição: {case.description}")

        if include_signature:
            doc.add_paragraph("Assinado digitalmente pelo perito responsável")

        doc.save(docx_path)
        return docx_path
    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {e}")
        return None
